from django.shortcuts import render,redirect, get_object_or_404,HttpResponseRedirect,HttpResponse
from django.views.generic import TemplateView, CreateView, ListView, DeleteView,DetailView
from django.urls import reverse_lazy
from django.utils.decorators import method_decorator
from django.contrib.auth.decorators import login_required
from authentication.decorators import user_is_instructor, user_is_student
import hashlib
from .forms import CourseCreateForm, AssignmentCreateForm, AssignmentSubmissionForm, AssignmentResultForm
from .models import Course, Assignment,AssignmentResult,AssignmentSubmission
from django.db import IntegrityError
from plag.plagcheck import handle_uploaded_file
import docx
from docx.shared import Inches
 
class HomeView(ListView):
    paginate_by = 6
    template_name = 'home.html'
    model = Course 
    context_object_name = 'course'

    def get_queryset(self):
        return self.model.objects.all().order_by('id')


class ExamView(TemplateView):
    template_name = 'core/exams.html'


# COURSE CREATE VIEW
class CourseCreateView(CreateView):
    template_name = 'core/instructor/course_create.html'
    form_class = CourseCreateForm
    extra_context = {
        'title': 'New Course'
    }
    success_url = reverse_lazy('core:course')

    @method_decorator(login_required(login_url=reverse_lazy('authentication:login')))
    def dispatch(self, request, *args, **kwargs):
        if not self.request.user.is_authenticated:
            return reverse_lazy('authentication:login')
        if self.request.user.is_authenticated and self.request.user.role != 'instructor':
            return reverse_lazy('authentication:login')
        return super().dispatch(self.request, *args, **kwargs)

    def form_valid(self, form):
        form.instance.user = self.request.user
        return super(CourseCreateView, self).form_valid(form)

    def post(self, request, *args, **kwargs):
        self.object = None
        form = self.get_form()
        if form.is_valid():
            return self.form_valid(form)
        else:
            return self.form_invalid(form)


# VIEW FOR COURSE LIST
class CourseView(ListView):
    model = Course
    template_name = 'core/instructor/courses.html'
    context_object_name = 'course'

    @method_decorator(login_required(login_url=reverse_lazy('authentication:login')))
    # @method_decorator(user_is_instructor, user_is_student)
    def dispatch(self, request, *args, **kwargs):
        return super().dispatch(self.request, *args, **kwargs)

    def get_queryset(self):
        return self.model.objects.all().order_by('-id')  # filter(user_id=self.request.user.id).order_by('-id')


def course_single(request, id):
    course = get_object_or_404(Course, id=id)
    return render(request, "core/instructor/view_course.html", {'course': course})


# ASSIGNMENT CREATE VIEW
class AssignmentCreateView(CreateView):
    template_name = 'core/instructor/assignment_create.html'
    form_class = AssignmentCreateForm
    extra_context = {
        'title': 'New Course'
    }
    success_url = reverse_lazy('core:assignment-list')

    @method_decorator(login_required(login_url=reverse_lazy('authentication:login')))
    def dispatch(self, request, *args, **kwargs):
        if not self.request.user.is_authenticated:
            return reverse_lazy('authentication:login')
        if self.request.user.is_authenticated and self.request.user.role != 'instructor':
            return reverse_lazy('authentication:login')
        return super().dispatch(self.request, *args, **kwargs)

    def form_valid(self, form):
        form.instance.user = self.request.user
        return super(AssignmentCreateView, self).form_valid(form)

    def post(self, request, *args, **kwargs):
        self.object = None
        form = self.get_form()
        if form.is_valid():
            return self.form_valid(form)
        else:
            return self.form_invalid(form)


# VIEW FOR ASSIGNMENT LIST 
class AssignmentView(ListView):
    model = Assignment
    template_name = 'core/instructor/assignments.html'
    context_object_name = 'assignment'

    @method_decorator(login_required(login_url=reverse_lazy('authentication:login')))
    # @method_decorator(user_is_student, user_is_instructor)
    def dispatch(self, request, *args, **kwargs):
        return super().dispatch(self.request, *args, **kwargs)

    def get_queryset(self):
        # AssignmentSubmission = AssignmentSubmission.objects.values('user_id').distinct()
        return self.model.objects.all()  # filter(user_id=self.request.user.id).order_by('-id')


class ResultlistView(ListView):
    model = AssignmentResult
    # AssignmentResult = AssignmentResult.objects.values('user_id').distinct()
    template_name = "core/instructor/display_result.html"
    context_object_name = 'assignmentresult'

    @method_decorator(login_required(login_url=reverse_lazy('authentication:login')))
    # @method_decorator(user_is_student, user_is_instructor)
    def dispatch(self, request, *args, **kwargs):
        return super().dispatch(self.request, *args, **kwargs)

    def get_queryset(self):
        # AssignmentSubmission = AssignmentSubmission.objects.values('user_id').distinct()
        return self.model.objects.all()  # filter(user_id=self.request.user.id).order_by('-id')

# result create view
class ResultCreateView(CreateView):
    template_name = 'core/instructor/assignment_feedback.html'
    form_class = AssignmentResultForm
    extra_context = {
        'title': 'result update'
    }
    success_url = reverse_lazy('core:assignment-submission-list')
    
    @method_decorator(login_required(login_url=reverse_lazy('authentication:login')))
    def dispatch(self, request, *args, **kwargs):
        print(self.request.user.id)
        # print(self.model.objects.values_list('id'))
        if not self.request.user.is_authenticated:
            return reverse_lazy('authentication:login')
        if self.request.user.is_authenticated and self.request.user.role != 'instructor':
            return reverse_lazy('authentication:login')
        return super().dispatch(self.request, *args, **kwargs)

    def form_valid(self, form):
        # try:
        #     AssignmentResult.objects.filter(id=self.request.user.id).delete()
        # except:
        #     print('cant delete')    
        form.instance.user = self.request.user
        return super(ResultCreateView, self).form_valid(form)

    def post(self, request, *args, **kwargs):
        self.object = None
        form = self.get_form()
        if form.is_valid():
            return self.form_valid(form)
        else:
            return self.form_invalid(form)    


def AssignmentDeleteView(request,pk,*args, **kwargs):
    AssignmentSubmission.objects.filter(id=pk).delete()
    # model = Assignment
            #success_url = reverse_lazy('core:assignment-list')
    return redirect('core:assignment-submission-list')

def PlagcheckView(request,pk,*args, **kwargs):
    print(AssignmentSubmission.objects.filter(id=pk).values('name'))#(id=pk))
    instance = AssignmentSubmission.objects.filter(id=pk).values('file')
    u_id = AssignmentSubmission.objects.filter(id=pk).values('university_id')
    name = AssignmentSubmission.objects.filter(id=pk).values('name')
    name=name[0]['name']
    u_id=u_id[0]['university_id']
    print('media/'+instance[0]['file'])
    f='media/'+instance[0]['file'] 
    try:
        filename=f"{name}.docx"
        document= docx.Document()
        document.add_heading("Plagerism Report")
        document.add_heading(f"Student Name : {name}")
        document.add_heading(f"University ID : {u_id}")
        document.add_heading("LINKS \n")
        ls=[]
        result=handle_uploaded_file(f)
        for key in result:
            if 'href' in key:
                ls=result['href']
                for k in ls: 
                    document.add_paragraph(k+'\n')   
            # print(k)
            else:
                document.add_heading("Plagerism : ")
                document.add_paragraph(result[key])    
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=Plagerism_Report_%s.docx'%filename
        document.save(response)

        return response
    except:
        print("The file uploaded by the student is corupted \n file should be in either of those \n .docx .pdf or .txt")
#     # if AssignmentSubmission.objects.FILES('file'):
#     #     print(ok)
#     # model = Assignment
#             #success_url = reverse_lazy('core:assignment-list')
    return redirect('core:assignment-submission-list')


class AssignmentSubmissionView(CreateView):
    template_name = 'core/instructor/assignment_submission.html'
    form_class = AssignmentSubmissionForm
    extra_context = {
        'title': 'New Exam'
    }
    success_url = reverse_lazy('core:home')
    @method_decorator(login_required(login_url=reverse_lazy('authentication:login')))
    def dispatch(self, request, *args, **kwargs):
        if not self.request.user.is_authenticated:
            return reverse_lazy('authentication:login')
        if self.request.user.is_authenticated and self.request.user.role != 'student':
            return reverse_lazy('authentication:login')
        return super().dispatch(self.request, *args, **kwargs)

    def form_valid(self, form):    
        form.instance.user = self.request.user
        return super(AssignmentSubmissionView, self).form_valid(form)
    def post(self, request, *args, **kwargs):
        print("here 2")
        self.object = None
        form = self.get_form()
        if form.is_valid():
            print("here 3")
            # handle_uploaded_file(request.FILES['file'])
            return self.form_valid(form)      
        else:
            print("here 4")
            return self.form_invalid(form)


# EXAM CREATE VIEW
# class ExamCreateView(CreateView):
#     template_name = 'core/instructor/exam_create.html'
#     form_class = ExamCreateForm
#     extra_context = {
#         'title': 'New Exam'
#     }
#     success_url = reverse_lazy('core:exam-list')

#     @method_decorator(login_required(login_url=reverse_lazy('authentication:login')))
#     def dispatch(self, request, *args, **kwargs):
#         if not self.request.user.is_authenticated:
#             return reverse_lazy('authentication:login')
#         if self.request.user.is_authenticated and self.request.user.role != 'instructor':
#             return reverse_lazy('authentication:login')
#         return super().dispatch(self.request, *args, **kwargs)

#     def form_valid(self, form):
#         form.instance.user = self.request.user
#         return super(ExamCreateView, self).form_valid(form)

#     def post(self, request, *args, **kwargs):
#         self.object = None
#         form = self.get_form()
#         if form.is_valid():
#             return self.form_valid(form)
#         else:
#             return self.form_invalid(form)


# # EXAM LIST VIEW
# class ExamListView(ListView):
#     model = Exam
#     template_name = 'core/instructor/exam_list.html'
#     context_object_name = 'exam'

#     @method_decorator(login_required(login_url=reverse_lazy('authentication:login')))
#     # @method_decorator(user_is_instructor)
#     def dispatch(self, request, *args, **kwargs):
#         return super().dispatch(self.request, *args, **kwargs)

#     def get_queryset(self):
#         return self.model.objects.all().order_by('-id')  # filter(user_id=self.request.user.id).order_by('-id')


# # EXAM DELETE VIEW
# class ExamDeleteView(DeleteView):
#     model = Exam
#     success_url = reverse_lazy('core:exam-list')


# # EXAM SUBMISSION VIEW
# class ExamSubmissionView(CreateView):
#     template_name = 'core/instructor/exam_submission.html'
#     form_class = ExamSubmissionForm
#     extra_context = {
#         'title': 'New Exam'
#     }
#     success_url = reverse_lazy('core:home')

#     @method_decorator(login_required(login_url=reverse_lazy('authentication:login')))
#     def dispatch(self, request, *args, **kwargs):
#         if not self.request.user.is_authenticated:
#             return reverse_lazy('authentication:login')
#         if self.request.user.is_authenticated and self.request.user.role != 'student':
#             return reverse_lazy('authentication:login')
#         return super().dispatch(self.request, *args, **kwargs)

#     def form_valid(self, form):
#         form.instance.user = self.request.user
#         return super(ExamSubmissionView, self).form_valid(form)

#     def post(self, request, *args, **kwargs):
#         self.object = None
#         form = self.get_form()
#         if form.is_valid():
#             return self.form_valid(form)
#         else:
#             return self.form_invalid(form)


# VIEW FOR Assignment Submission List
class AssignmentSubmissionListView(ListView):
    model = AssignmentSubmission
    template_name = 'core/instructor/assignment_submission_list.html'
    context_object_name = 'assignment_submission'

    @method_decorator(login_required(login_url=reverse_lazy('authentication:login')))
    # @method_decorator(user_is_instructor, user_is_student)
    def dispatch(self, request, *args, **kwargs):
        return super().dispatch(self.request, *args, **kwargs)

    def get_queryset(self):
        return self.model.objects.all().order_by('id')  # filter(user_id=self.request.user.id).order_by('-id')


# EXAM DELETE VIEW
class AssignmentSubmissionDelete(DeleteView):
    model = AssignmentSubmission
    success_url = reverse_lazy('core:assignment-submission-list')


# # VIEW FOR Assignment Submission List
# class ExamSubmissionListView(ListView):
#     model = ExamSubmission
#     template_name = 'core/instructor/exam_submission_list.html'
#     context_object_name = 'exam_submission'

#     @method_decorator(login_required(login_url=reverse_lazy('authentication:login')))
#     # @method_decorator(user_is_instructor, user_is_student)
#     def dispatch(self, request, *args, **kwargs):
#         return super().dispatch(self.request, *args, **kwargs)

#     def get_queryset(self):
#         return self.model.objects.all().order_by('-id')  # filter(user_id=self.request.user.id).order_by('-id')


# class ExamSubmissionDelete(DeleteView):
#     model = ExamSubmission
#     success_url = reverse_lazy('core:exam-submission-list')
